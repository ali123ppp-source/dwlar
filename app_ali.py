import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from streamlit_gsheets import GSheetsConnection
from xhtml2pdf import pisa  # المكتبة الاحترافية لإنتاج ملف PDF حقيقي ومباشر للهاتف

# =============================================================================
# 🔐 إعدادات الهوية وقفل البرنامج (تغيير اسم الوكيل ورابط الأيقونة هنا)
# =============================================================================
AUTHORIZED_AGENT_NAME = "علي صباح حسن"

# 🖼️ رابط أيقونة التطبيق لشاشة الهاتف (تظهر عند الحفظ والتثبيت على شاشة الهاتف الرئيسية)
APP_LOGO_URL = "https://cdn-icons-png.flaticon.com/512/5606/5606132.png" 

# -----------------------------------------------------------------------------
# 1. إعدادات النظام والمظهر والأيقونات للهاتف
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title=f"نظام الوكيل - {AUTHORIZED_AGENT_NAME}", 
    page_icon=APP_LOGO_URL, 
    layout="wide", 
    initial_sidebar_state="collapsed"
)

# حقن الأيقونة لتظهر كأيقونة تطبيق رسمي عند التثبيت على شاشة الآيفون والأندرويد
st.markdown(f'<link rel="apple-touch-icon" href="{APP_LOGO_URL}"><link rel="icon" type="image/png" href="{APP_LOGO_URL}">', unsafe_allow_html=True)

# الـ CSS المتقدم والآمن لتصميم أزرار زرقاء ملكية واضحة جداً للهواتف والأجهزة الذكية
st.markdown("""
    <style>
    th, td { text-align: right !important; direction: rtl !important; }
    
    /* 🔵 أزرار زرقاء ملكية بارزة وواضحة جداً وتصميم مريح للنقر بالهاتف */
    div.stButton > button, div.stDownloadButton > button { 
        background-color: #005C99 !important; 
        color: #ffffff !important; 
        width: 100%; 
        font-weight: bold !important; 
        font-size: 16px !important;
        border-radius: 8px !important; 
        border: none !important; 
        height: 52px !important; 
        box-shadow: 0px 4px 8px rgba(0, 0, 0, 0.2) !important; 
        transition: 0.3s ease;
        margin-bottom: 10px;
    }
    
    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #004473 !important;
        box-shadow: 0px 2px 4px rgba(0, 0, 0, 0.1) !important;
    }

    .report-box { 
        background-color: var(--secondary-background-color); color: var(--text-color); padding: 15px; border-radius: 8px; 
        border-right: 5px solid #005C99; text-align: right; margin-bottom: 15px; box-shadow: 0px 2px 5px rgba(0,0,0,0.05);
    }
    .report-box h4 { color: var(--text-color); margin: 5px 0; }
    </style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. الاتصال بجوجل شيت (الأرشيف)
# -----------------------------------------------------------------------------
try:
    conn = st.connection("gsheets", type=GSheetsConnection)
except Exception as e:
    st.error("تنبيه: لم يتم ضبط إعدادات الاتصال بجوجل شيت بعد.")

# -----------------------------------------------------------------------------
# 3. محرك المعالجة والمقارنة ووظائف التقارير
# -----------------------------------------------------------------------------
def extract_clean_records(file_obj):
    doc = Document(file_obj)
    records = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells):
                continue
            name_idx = -1
            max_len = 0
            for i, c in enumerate(cells):
                if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                    if len(c) > max_len: max_len = len(c); name_idx = i
            if name_idx == -1: continue
            card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
            if not card_indices: continue
            card_num = cells[card_indices[1]] if len(card_indices) >= 2 else cells[card_indices[0]]
            seq = "-"
            for i in range(len(cells)-1, card_indices[-1], -1):
                if cells[i].isdigit(): seq = cells[i]; break
            digit_cells = [int(cells[i]) for i in range(name_idx) if cells[i].isdigit()]
            if len(digit_cells) >= 3: withheld, eligible, total = digit_cells[0], digit_cells[1], digit_cells[2]
            elif len(digit_cells) == 2: withheld, eligible, total = 0, digit_cells[0], digit_cells[1]
            else: continue
            records[card_num] = {"seq": seq, "name": cells[name_idx], "total": total, "eligible": eligible, "withheld": withheld}
    return records

def compare_records(old_data, new_data):
    results = []
    counters = {
        "total_fam": 0, "eligible_fam": 0, "withheld_fam": 0, "added_fam": 0, "deleted_fam": 0,
        "net_total": 0, "net_eligible": 0, "net_withheld": 0
    }
    all_cards = set(old_data.keys()).union(set(new_data.keys()))
    for card in all_cards:
        if card in old_data and card in new_data:
            old_v, new_v = old_data[card], new_data[card]
            diff_total = old_v["total"] != new_v["total"]
            diff_elig = old_v["eligible"] != new_v["eligible"]
            diff_with = old_v["withheld"] != new_v["withheld"]
            if diff_total or diff_elig or diff_with:
                if diff_total: counters["total_fam"] += 1; counters["net_total"] += (new_v["total"] - old_v["total"])
                if diff_elig: counters["eligible_fam"] += 1; counters["net_eligible"] += (new_v["eligible"] - old_v["eligible"])
                if diff_with: counters["withheld_fam"] += 1; counters["net_withheld"] += (new_v["withheld"] - old_v["withheld"])
                results.append({
                    "التسلسل": new_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": old_v["name"], "الاسم (حالياً)": new_v["name"],
                    "الكلية (سابقاً)": old_v["total"], "الكلية (حالياً)": new_v["total"], "المستحقة (سابقاً)": old_v["eligible"], "المستحقة (حالياً)": new_v["eligible"],
                    "المحجوبين (سابقاً)": old_v["withheld"], "المحجوبين (حالياً)": new_v["withheld"]
                })
        elif card in old_data and card not in new_data:
            old_v = old_data[card]
            counters["deleted_fam"] += 1
            counters["net_total"] -= old_v["total"]; counters["net_eligible"] -= old_v["eligible"]; counters["net_withheld"] -= old_v["withheld"]
            results.append({
                "التسلسل": old_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": old_v["name"], "الاسم (حالياً)": "❌ (محذوف / منقول)",
                "الكلية (سابقاً)": old_v["total"], "الكلية (حالياً)": 0, "المستحقة (سابقاً)": old_v["eligible"], "المستحقة (حالياً)": 0,
                "المحجوبين (سابقاً)": old_v["withheld"], "المحجوبين (حالياً)": 0
            })
        elif card not in old_data and card in new_data:
            new_v = new_data[card]
            counters["added_fam"] += 1
            counters["net_total"] += new_v["total"]; counters["net_eligible"] += new_v["eligible"]; counters["net_withheld"] += new_v["withheld"]
            results.append({
                "التسلسل": new_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": "✨ (مضاف حديثاً)", "الاسم (حالياً)": new_v["name"],
                "الكلية (سابقاً)": 0, "الكلية (حالياً)": new_v["total"], "المستحقة (سابقاً)": 0, "المستحقة (حالياً)": new_v["eligible"],
                "المحجوبين (سابقاً)": 0, "المحجوبين (حالياً)": new_v["withheld"]
            })
    return results, counters

def create_word_table_report(df, title):
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cols = list(df.columns)[::-1]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = str(row[col])
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 📄 توليد ملف الـ PDF الاحترافي المباشر باللغة العربية والخطوط الهندسية الواضحة للأرقام والأسماء
def create_pdf_file_report(df, title, agent_name):
    html_table = df.to_html(index=False)
    html_content = f"""
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Amiri&display=swap');
            @page {{ size: a4 landscape; margin: 1cm; }}
            body {{ font-family: 'Amiri', serif; direction: rtl; text-align: right; color: #333; }}
            .header {{ text-align: center; margin-bottom: 20px; border-bottom: 2px solid #005C99; padding-bottom: 10px; }}
            .header h1 {{ color: #005C99; font-size: 22px; margin: 0; }}
            .header p {{ font-size: 12px; color: #555; margin: 4px 0; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 10px; direction: rtl; }}
            th, td {{ border: 1px solid #999; padding: 6px; text-align: center; font-size: 11px; }}
            th {{ background-color: #005C99; color: white; }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{title}</h1>
            <p>الوكيل المعتمد: {agent_name} | تاريخ الاستخراج: {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
        </div>
        {html_table}
    </body>
    </html>
    """
    pdf_buffer = BytesIO()
    pisa.CreatePDF(BytesIO(html_content.encode("utf-8")), pdf_buffer)
    pdf_buffer.seek(0)
    return pdf_buffer

# -----------------------------------------------------------------------------
# 4. واجهة التطبيق الرئيسية 
# -----------------------------------------------------------------------------
st.markdown(f"<h2 style='text-align: right;'>نظام المقارنة الذكي - نسخة الوكيل: {AUTHORIZED_AGENT_NAME} 📄</h2>", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🔎 إجراء مقارنة جديدة", "📜 الأرشيف التاريخي"])

with tab1:
    col1, col2 = st.columns(2)
    with col1: new_file = st.file_uploader("الملف الجديد (docx)", key="n_f")
    with col2: old_file = st.file_uploader("الملف القديم (docx)", key="o_f")

    if st.button("🚀 تشغيل الفحص السحابي وبدء المقارنة"):
        if old_file and new_file:
            if AUTHORIZED_AGENT_NAME not in old_file.name or AUTHORIZED_AGENT_NAME not in new_file.name:
                st.error(f"❌ عذراً! هذه النسخة مقفلة ومخصصة حصراً لملفات الوكيل ({AUTHORIZED_AGENT_NAME}). تأكد من أن اسم الوكيل مكتوب بوضوح في اسم الملف المرفوع.")
            else:
                with st.spinner('جاري التحليل واستخراج المتغيرات...'):
                    old_data = extract_clean_records(old_file)
                    new_data = extract_clean_records(new_file)
                    results, counters = compare_records(old_data, new_data)
                    
                    if results:
                        df_results = pd.DataFrame(results)[["التسلسل", "رقم البطاقة", "الاسم (سابقاً)", "الاسم (حالياً)", "الكلية (سابقاً)", "الكلية (حالياً)", "المستحقة (سابقاً)", "المستحقة (حالياً)", "المحجوبين (سابقاً)", "المحجوبين (حالياً)"]]
                        st.session_state['c_results'] = df_results
                        st.session_state['c_counters'] = counters
                        st.session_state['c_filename'] = new_file.name.rsplit('.', 1)[0]
                        st.session_state['show_table'] = True
                        st.success("✅ تم الفحص بنجاح والمحافظة على الجدول!")
                    else: 
                        st.session_state['show_table'] = False
                        st.info("تطابق كامل بين الملفين، لا توجد فروقات.")
        else: 
            st.warning("يرجى رفع الملفات أولاً.")

    # -----------------------------------------------------------------------------
    # 🌟 عرض واستمرار جدول النتائج في الجلسة الحالية
    # -----------------------------------------------------------------------------
    if st.session_state.get('show_table', False) and 'c_results' in st.session_state:
        df_res = st.session_state['c_results']
        cnt = st.session_state['c_counters']
        
        st.markdown("<h4 style='text-align: right;'>📊 خلاصة المتغيرات للعملية الحالية</h4>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1: st.markdown(f"<div class='report-box'>🔹 حركة الكلية<br><h4>{cnt['total_fam']} عائلة</h4>الصافي: {cnt['net_total']:+d}</div>", unsafe_allow_html=True)
        with c2: st.markdown(f"<div class='report-box'>🔹 حركة المستحقة<br><h4>{cnt['eligible_fam']} عائلة</h4>الصافي: {cnt['net_eligible']:+d}</div>", unsafe_allow_html=True)
        with c3: st.markdown(f"<div class='report-box'>🔹 الحالات<br><h4>مضاف: {cnt['added_fam']} | محذوف: {cnt['deleted_fam']}</h4></div>", unsafe_allow_html=True)
        
        if st.button("💾 ترحيل وحفظ هذه العملية في الأرشيف"):
            with st.spinner("جاري الترحيل..."):
                try:
                    existing_df = conn.read(worksheet="Sheet1", ttl=0)
                    new_row = pd.DataFrame([{
                        "التاريخ": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "الوكيل / المستخدم": AUTHORIZED_AGENT_NAME,
                        "اسم الملف المفحوص": st.session_state['c_filename'],
                        "حركة الكلية (عائلات)": cnt['total_fam'],
                        "صافي الأفراد (كلية)": cnt['net_total'],
                        "حركة المستحقة (عائلات)": cnt['eligible_fam'],
                        "صافي الأفراد (مستحقة)": cnt['net_eligible'],
                        "العائلات المضافة": cnt['added_fam'],
                        "العائلات المحذوفة": cnt['deleted_fam']
                    }])
                    updated_df = pd.concat([existing_df, new_row], ignore_index=True)
                    conn.update(worksheet="Sheet1", data=updated_df)
                    st.success("🚀 تم ترحيل البيانات بنجاح إلى الأرشيف السحابي واحتفظنا بجدولك!")
                except Exception as ex:
                    st.error(f"فشل الاتصال بالأرشيف: {ex}")

        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown("<h5 style='text-align: right;'>🎛️ تصفية وعرض فئات معينة من الجدول:</h5>", unsafe_allow_html=True)
        
        filter_option = st.radio(
            "اختر الفئة المراد عرضها وتحميلها:",
            ["الكل (جميع المتغيرات)", "✨ المضافين فقط", "❌ المحذوفين فقط", "🔄 تعديلات الأفراد فقط"],
            horizontal=True,
            label_visibility="collapsed"
        )
        
        df_filtered = df_res.copy()
        if filter_option == "✨ المضافين فقط":
            df_filtered = df_filtered[df_filtered["الاسم (سابقاً)"].str.contains("مضاف", na=False)]
        elif filter_option == "❌ المحذوفين فقط":
            df_filtered = df_filtered[df_filtered["الاسم (حالياً)"].str.contains("❌|محذوف", na=False)]
        elif filter_option == "🔄 تعديلات الأفراد فقط":
            df_filtered = df_filtered[~df_filtered["الاسم (سابقاً)"].str.contains("مضاف", na=False) & ~df_filtered["الاسم (حالياً)"].str.contains("❌|محذوف", na=False)]
            
        st.dataframe(df_filtered, use_container_width=True, hide_index=True)
        
        report_title = f"تقرير المتغيرات ({filter_option.split()[0]})"
        
        # أزرار التحميل المباشرة والاحترافية بجانب بعضهما البعض
        btn_col1, btn_col2 = st.columns(2)
        
        with btn_col1:
            word_report = create_word_table_report(df_filtered, f"{report_title} - {st.session_state['c_filename']}")
            st.download_button(
                label="📥 تحميل الجدول الحالي كتقرير (Word)",
                data=word_report,
                file_name=f"تقرير_{filter_option.split()[0]}_{st.session_state['c_filename']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            
        with btn_col2:
            # 📁 توليد وتحميل ملف PDF حقيقي ومباشر إلى ذاكرة الهاتف بشكل فوري
            pdf_report = create_pdf_file_report(df_filtered, f"{report_title} - {st.session_state['c_filename']}", AUTHORIZED_AGENT_NAME)
            st.download_button(
                label="📄 تحميل الجدول الحالي كملف (PDF) مباشر",
                data=pdf_report,
                file_name=f"تقرير_{filter_option.split()[0]}_{st.session_state['c_filename']}.pdf",
                mime="application/pdf",
            )

with tab2:
    st.markdown("<h3 style='text-align: right;'>📜 العمليات المؤرشفة</h3>", unsafe_allow_html=True)
    if st.button("🔄 تحديث وجلب البيانات الحالية من الأرشيف"):
        try:
            archive_df = conn.read(worksheet="Sheet1", ttl=0)
            if not archive_df.empty:
                agent_archive = archive_df[archive_df["الوكيل / المستخدم"] == AUTHORIZED_AGENT_NAME]
                if not agent_archive.empty:
                    st.dataframe(agent_archive, use_container_width=True, hide_index=True)
                else:
                    st.info("لا توجد عمليات سابقة مؤرشفة باسمك حالياً.")
            else:
                st.info("الأرشيف الكلي فارغ حالياً.")
        except Exception as ex:
            st.error(f"لم نتمكن من قراءة الأرشيف: {ex}")
