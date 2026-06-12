import streamlit as st
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
# استيراد مكتبة الربط مع جوجل شيت
from streamlit_gsheets import GSheetsConnection

# -----------------------------------------------------------------------------
# 1. إعدادات النظام والمظهر
# -----------------------------------------------------------------------------
st.set_page_config(page_title="نظام المقارنة السحابي المطور", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { 
        background-color: var(--primary-color); color: white; width: 100%; font-weight: bold; border-radius: 8px; border: none; height: 45px;
    }
    .report-box { 
        background-color: var(--secondary-background-color); color: var(--text-color); padding: 15px; border-radius: 8px; 
        border-right: 5px solid var(--primary-color); text-align: right; margin-bottom: 15px; box-shadow: 0px 2px 5px rgba(0,0,0,0.05);
    }
    .report-box h4 { color: var(--text-color); margin: 5px 0; }
    </style>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. إنشاء الاتصال بجوجل شيت
# -----------------------------------------------------------------------------
try:
    conn = st.connection("gsheets", type=GSheetsConnection)
except Exception as e:
    st.error("تنبيه: لم يتم ضبط إعدادات الاتصال بجوجل شيت بعد في ملف Secrets.")

# نظام الحسابات الافتراضي
USER_CREDENTIALS = {
    "agent_hillah": "pass1234",
    "agent_baghdad": "baghdad2026",
    "admin": "master_root_99"
}

# -----------------------------------------------------------------------------
# 3. محرك المعالجة والمقارنة (بايثون)
# -----------------------------------------------------------------------------
def extract_clean_records(file_obj):
    doc = Document(file_obj)
    records = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells):
                continue
            name_idx = -1
            max_len = 0
            for i, c in enumerate(cells):
                if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                    if len(c) > max_len: max_len = len(c); name_idx = i
            if name_idx == -1: continue
            card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
            if not card_indices: continue
            card_num = cells[card_indices[1]] if len(card_indices) >= 2 else cells[card_indices[0]]
            seq = "-"
            for i in range(len(cells)-1, card_indices[-1], -1):
                if cells[i].isdigit(): seq = cells[i]; break
            digit_cells = [int(cells[i]) for i in range(name_idx) if cells[i].isdigit()]
            if len(digit_cells) >= 3: withheld, eligible, total = digit_cells[0], digit_cells[1], digit_cells[2]
            elif len(digit_cells) == 2: withheld, eligible, total = 0, digit_cells[0], digit_cells[1]
            else: continue
            records[card_num] = {"seq": seq, "name": cells[name_idx], "total": total, "eligible": eligible, "withheld": withheld}
    return records

def compare_records(old_data, new_data):
    results = []
    counters = {
        "total_fam": 0, "eligible_fam": 0, "withheld_fam": 0, "added_fam": 0, "deleted_fam": 0,
        "net_total": 0, "net_eligible": 0, "net_withheld": 0
    }
    all_cards = set(old_data.keys()).union(set(new_data.keys()))
    for card in all_cards:
        if card in old_data and card in new_data:
            old_v, new_v = old_data[card], new_data[card]
            diff_total = old_v["total"] != new_v["total"]
            diff_elig = old_v["eligible"] != new_v["eligible"]
            diff_with = old_v["withheld"] != new_v["withheld"]
            if diff_total or diff_elig or diff_with:
                if diff_total: counters["total_fam"] += 1; counters["net_total"] += (new_v["total"] - old_v["total"])
                if diff_elig: counters["eligible_fam"] += 1; counters["net_eligible"] += (new_v["eligible"] - old_v["eligible"])
                if diff_with: counters["withheld_fam"] += 1; counters["net_withheld"] += (new_v["withheld"] - old_v["withheld"])
                results.append({
                    "التسلسل": new_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": old_v["name"], "الاسم (حالياً)": new_v["name"],
                    "الكلية (سابقاً)": old_v["total"], "الكلية (حالياً)": new_v["total"], "المستحقة (سابقاً)": old_v["eligible"], "المستحقة (حالياً)": new_v["eligible"],
                    "المحجوبين (سابقاً)": old_v["withheld"], "المحجوبين (حالياً)": new_v["withheld"]
                })
        elif card in old_data and card not in new_data:
            old_v = old_data[card]
            counters["deleted_fam"] += 1
            counters["net_total"] -= old_v["total"]; counters["net_eligible"] -= old_v["eligible"]; counters["net_withheld"] -= old_v["withheld"]
            results.append({
                "التسلسل": old_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": old_v["name"], "الاسم (حالياً)": "❌ (محذوف / منقول)",
                "الكلية (سابقاً)": old_v["total"], "الكلية (حالياً)": 0, "المستحقة (سابقاً)": old_v["eligible"], "المستحقة (حالياً)": 0,
                "المحجوبين (سابقاً)": old_v["withheld"], "المحجوبين (حالياً)": 0
            })
        elif card not in old_data and card in new_data:
            new_v = new_data[card]
            counters["added_fam"] += 1
            counters["net_total"] += new_v["total"]; counters["net_eligible"] += new_v["eligible"]; counters["net_withheld"] += new_v["withheld"]
            results.append({
                "التسلسل": new_v["seq"], "رقم البطاقة": card, "الاسم (سابقاً)": "✨ (مضاف حديثاً)", "الاسم (حالياً)": new_v["name"],
                "الكلية (سابقاً)": 0, "الكلية (حالياً)": new_v["total"], "المستحقة (سابقاً)": 0, "المستحقة (حالياً)": new_v["eligible"],
                "المحجوبين (سابقاً)": 0, "المحجوبين (حالياً)": new_v["withheld"]
            })
    return results, counters

def create_word_table_report(df, title):
    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cols = list(df.columns)[::-1]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = str(col)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = str(row[col])
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# 4. بوابه تسجيل الدخول
# -----------------------------------------------------------------------------
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    st.markdown("<h2 style='text-align: right;'>🔐 تسجيل الدخول للنظام السحابي</h2>", unsafe_allow_html=True)
    u_in = st.text_input("👤 اسم المستخدم:")
    p_in = st.text_input("🔑 كلمة المرور:", type="password")
    if st.button("تسجيل الدخول"):
        if u_in in USER_CREDENTIALS and USER_CREDENTIALS[u_in] == p_in:
            st.session_state['logged_in'] = True
            st.session_state['username'] = u_in
            st.rerun()
        else:
            st.error("بيانات الدخول خاطئة.")
    st.stop()

# -----------------------------------------------------------------------------
# 5. واجهة التطبيق الرئيسية (تبويبات)
# -----------------------------------------------------------------------------
st.markdown(f"<p style='text-align: left; color: gray;'>المستخدم: <b>{st.session_state['username']}</b></p>", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🔎 إجراء مقارنة جديدة", "📜 الأرشيف التاريخي (جوجل شيت)"])

with tab1:
    st.markdown("<h3 style='text-align: right;'>لوحة المطابقة الرقمية</h3>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1: new_file = st.file_uploader("الملف الجديد (docx)", key="n_f")
    with col2: old_file = st.file_uploader("الملف القديم (docx)", key="o_f")

    if st.button("تشغيل الفحص السحابي"):
        if old_file and new_file:
            with st.spinner('جاري التحليل...'):
                old_data = extract_clean_records(old_file)
                new_data = extract_clean_records(new_file)
                results, counters = compare_records(old_data, new_data)
                if results:
                    df_results = pd.DataFrame(results)[["التسلسل", "رقم البطاقة", "الاسم (سابقاً)", "الاسم (حالياً)", "الكلية (سابقاً)", "الكلية (حالياً)", "المستحقة (سابقاً)", "المستحقة (حالياً)", "المحجوبين (سابقاً)", "المحجوبين (حالياً)"]]
                    st.session_state['c_results'] = df_results
                    st.session_state['c_counters'] = counters
                    st.session_state['c_filename'] = new_file.name.rsplit('.', 1)[0]
                    st.success("✅ تم الفحص!")
                else: st.info("تطابق كامل.")
        else: st.warning("يرجى رفع الملفات.")

    # عرض النتائج وزر الحفظ
    if 'c_results' in st.session_state:
        df_res = st.session_state['c_results']
        cnt = st.session_state['c_counters']
        
        st.markdown("<h4 style='text-align: right;'>📊 خلاصة المتغيرات الحالية</h4>", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1: st.markdown(f"<div class='report-box'>🔹 حركة الكلية<br><h4>{cnt['total_fam']} عائلة</h4>الصافي: {cnt['net_total']:+d}</div>", unsafe_allow_html=True)
        with c2: st.markdown(f"<div class='report-box'>🔹 حركة المستحقة<br><h4>{cnt['eligible_fam']} عائلة</h4>الصافي: {cnt['net_eligible']:+d}</div>", unsafe_allow_html=True)
        with c3: st.markdown(f"<div class='report-box'>🔹 الحالات<br><h4>مضاف: {cnt['added_fam']} | محذوف: {cnt['deleted_fam']}</h4></div>", unsafe_allow_html=True)
        
        # 💾 زر حفظ النتائج في جوجل شيت
        if st.button("💾 ترحيل وحفظ هذه العملية في أرشيف جوجل شيت السحابي"):
            with st.spinner("جاري الترحيل للسحابة..."):
                try:
                    # 1. قراءة البيانات الحالية من الشيت
                    existing_df = conn.read(worksheet="Sheet1", ttl=0)
                    # 2. تجهيز السطر الجديد
                    new_row = pd.DataFrame([{
                        "التاريخ": datetime.now().strftime("%Y-%m-%d %H:%M"),
                        "الوكيل / المستخدم": st.session_state['username'],
                        "اسم الملف المفحوص": st.session_state['c_filename'],
                        "حركة الكلية (عائلات)": cnt['total_fam'],
                        "صافي الأفراد (كلية)": cnt['net_total'],
                        "حركة المستحقة (عائلات)": cnt['eligible_fam'],
                        "صافي الأفراد (مستحقة)": cnt['net_eligible'],
                        "العائلات المضافة": cnt['added_fam'],
                        "العائلات المحذوفة": cnt['deleted_fam']
                    }])
                    # 3. دمج وتحديث جوجل شيت
                    updated_df = pd.concat([existing_df, new_row], ignore_index=True)
                    conn.update(worksheet="Sheet1", data=updated_df)
                    st.success("🚀 تم ترحيل البيانات بنجاح إلى جوجل شيت وتحديث الأرشيف!")
                except Exception as ex:
                    st.error(f"فشل الاتصال بالشيت، تأكد من الإعدادات: {ex}")

        st.dataframe(df_res, use_container_width=True, hide_index=True)

with tab2:
    st.markdown("<h3 style='text-align: right;'>📜 العمليات المؤرشفة في جوجل شيت</h3>", unsafe_allow_html=True)
    if st.button("🔄 تحديث وجلب البيانات الحالية من جوجل شيت"):
        try:
            archive_df = conn.read(worksheet="Sheet1", ttl=0)
            if not archive_df.empty:
                st.dataframe(archive_df, use_container_width=True, hide_index=True)
            else:
                st.info("الأرشيف فارغ حالياً.")
        except Exception as ex:
            st.error(f"لم نتمكن من قراءة الشيت: {ex}")