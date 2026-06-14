import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from streamlit_gsheets import GSheetsConnection
import datetime

# -----------------------------------------------------------------------------
# 1. إعدادات النظام والتصميم الديناميكي المطور والفاخر للواجهة
# -----------------------------------------------------------------------------
st.set_page_config(page_title="نظام التدقيق الشامل لبيانات الوكلاء", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { background-color: #1B5E20; color: white; width: 100%; font-weight: bold; border-radius: 8px; border: none; height: 48px; font-size: 16px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
    div.stButton > button:hover { background-color: #2E7D32; }
    
    .report-box { background-color: var(--secondary-background-color); color: var(--text-color); padding: 15px; border-radius: 10px; border-right: 6px solid #1B5E20; text-align: right; margin-bottom: 12px; font-weight: bold; box-shadow: 0 2px 5px rgba(0,0,0,0.05); }
    
    .family-card-wrapper { padding: 12px; border-radius: 8px; margin-bottom: 10px; border-right: 6px solid #9e9e9e; background-color: var(--background-color); }
    .card-matched { border-right-color: #2E7D32 !important; background-color: rgba(46, 125, 50, 0.03); }
    .card-modified { border-right-color: #F57C00 !important; background-color: rgba(245, 124, 0, 0.03); }
    .card-added { border-right-color: #0288D1 !important; background-color: rgba(2, 136, 209, 0.03); }
    .card-deleted { border-right-color: #C62828 !important; background-color: rgba(198, 40, 40, 0.03); }
    
    .metric-grid-box { background: var(--secondary-background-color); padding: 10px 14px; border-radius: 8px; border: 1px solid rgba(0,0,0,0.08); text-align: center; margin-bottom: 8px; }
    .metric-title { font-size: 12px; color: gray; margin-bottom: 4px; display: block; font-weight: normal; }
    .metric-value { font-size: 18px; font-weight: bold; display: block; }
    
    .color-total { color: #0288D1; }
    .color-eligible { color: #2E7D32; }
    .color-withheld { color: #C62828; }
    </style>
""", unsafe_allow_html=True)

# الاتصال بقوقل شيت للأرشفة التاريخية
conn = None
try:
    conn = st.connection("gsheets", type=GSheetsConnection)
except Exception:
    st.error("⚠️ فشل الاتصال بقاعدة البيانات السحابية (Google Sheets)، تحقق من الإعدادات.")

# -----------------------------------------------------------------------------
# 2. محرك الذاكرة المستمرة للمتصفح (حماية ضد عمل Refresh)
# -----------------------------------------------------------------------------
@st.cache_resource
def get_persistent_memory():
    return {"df": None, "counters": None, "filename": None}

browser_memory = get_persistent_memory()

if 'c_results' not in st.session_state and browser_memory["df"] is not None:
    st.session_state['c_results'] = browser_memory["df"]
    st.session_state['c_counters'] = browser_memory["counters"]
    st.session_state['c_filename'] = browser_memory["filename"]
    st.toast("🔄 تم استعادة سجلات الوكيل بالكامل تلقائياً من ذاكرة المتصفح!", icon="⚡")

# -----------------------------------------------------------------------------
# 3. محرك تحليل السجلات واستخراج الأسماء الرباعية والتسلسل
# -----------------------------------------------------------------------------
def parse_row(cells):
    clean_cells = [str(c).strip().replace('\n', ' ') if c is not None else "" for c in cells]
    joined = "".join(clean_cells)
    
    if not any(clean_cells) or "المركز" in joined or "الوكيل" in joined or "اسم" in joined or "زكرمل" in joined: 
        return {}
        
    name_idx, max_len = -1, 0
    for i, c in enumerate(clean_cells):
        if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
            if len(c) > max_len: max_len, name_idx = len(c), i
                
    if name_idx == -1: return {}
    
    final_quad_name = clean_cells[name_idx]

    card_cands = [c for c in clean_cells if c.isdigit() and len(c) >= 5]
    card_num = card_cands[0] if card_cands else "-"
    if card_num == "-": return {}

    small_before = [int(clean_cells[i]) for i in range(name_idx) if clean_cells[i].isdigit() and len(clean_cells[i]) < 5]
    small_after = [int(clean_cells[i]) for i in range(name_idx + 1, len(clean_cells)) if clean_cells[i].isdigit() and len(clean_cells[i]) < 5]
    
    seq, total, eligible, withheld = "-", 0, 0, 0
    
    if len(small_before) >= 2:
        if len(small_before) >= 3: withheld, eligible, total = small_before[-3], small_before[-2], small_before[-1]
        else: withheld, eligible, total = 0, small_before[-2], small_before[-1]
        seq = small_after[0] if small_after else "-"
    elif len(small_after) >= 2:
        if len(small_after) >= 3: total, eligible, withheld = small_after[0], small_after[1], small_after[2]
        else: total, eligible, withheld = small_after[0], small_after[1], 0
        seq = small_before[0] if small_before else "-"
    else: return {}
        
    return {card_num: {"seq": seq, "name": final_quad_name, "total": total, "eligible": eligible, "withheld": withheld}}

def extract_clean_records(file_obj):
    records = {}
    doc = Document(file_obj)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            records.update(parse_row(cells))
    return records

# -----------------------------------------------------------------------------
# 4. كود المقارنة الرصين والمطور (محرك الجرد المحصن والمحاسبي المتقدم)
# -----------------------------------------------------------------------------
def compare_records(old_data, new_data):
    results = []
    counters = {
        "added_fam": 0, "deleted_fam": 0, "unchanged_fam": 0, "modified_fam": 0,
        "net_total": 0, "net_eligible": 0, "net_withheld": 0, "data_errors": 0
    }
    
    all_cards = set(old_data.keys()).union(set(new_data.keys()))
    
    for card in sorted(all_cards):
        # دالة حماية داخلية لتحويل آمن للحقول التموينية إلى أرقام صلبة
        def safe_ints(data_dict):
            try: t = int(data_dict.get("total", 0))
            except: t = 0
            try: e = int(data_dict.get("eligible", 0))
            except: e = 0
            try: w = int(data_dict.get("withheld", 0))
            except: w = 0
            return t, e, w

        # الحالة الأولى: العائلة موجودة في الكشفين (فحص تطابق وتعديل الحقول الداخلية)
        if card in old_data and card in new_data:
            old_v, new_v = old_data[card], new_data[card]
            old_t, old_e, old_w = safe_ints(old_v)
            new_t, new_e, new_w = safe_ints(new_v)
            
            old_name = str(old_v.get("name", "")).strip()
            new_name = str(new_v.get("name", "")).strip()
            seq = new_v.get("seq", "-")
            
            # طبقة تدقيق الرصانة: الفحص الرياضي للمعادلة (الكلية = المستحقة + المحجوبين)
            audit_note = ""
            if old_t != (old_e + old_w):
                audit_note += " [⚠️ خلل حسابي بالملف السابق]"
            if new_t != (new_e + new_w):
                audit_note += " [⚠️ خلل حسابي بالملف الحالي]"
            if audit_note:
                counters["data_errors"] += 1

            # جرد تفصيلي دقيق لكل حقل على حدة
            chg_total = old_t != new_t
            chg_eligible = old_e != new_e
            chg_withheld = old_w != new_w
            chg_name = old_name != new_name
            
            if chg_total or chg_eligible or chg_withheld or chg_name:
                counters["modified_fam"] += 1
                counters["net_total"] += (new_t - old_t)
                counters["net_eligible"] += (new_e - old_e)
                counters["net_withheld"] += (new_w - old_w)
                
                # بناء وصف التغيير الرصين بناءً على الحقل المتغير بدقة
                detected_changes = []
                if chg_name: detected_changes.append("اسم رب الأسرة")
                if chg_total: detected_changes.append("الكلية")
                if chg_eligible: detected_changes.append("المستحقة")
                if chg_withheld: detected_changes.append("المحجوبين")
                
                status_text = f"🟡 قيد معدل: تغيير في ({'، '.join(detected_changes)})" + audit_note
                
                results.append({
                    "التسلسل": seq, "رقم البطاقة": card,
                    "الاسم الرباعي (سابقاً)": old_name, "الاسم الرباعي (حالياً)": new_name,
                    "الكلية (سابقاً)": old_t, "الكلية (حالياً)": new_t,
                    "المستحقة (سابقاً)": old_e, "المستحقة (حالياً)": new_e,
                    "المحجوبين (سابقاً)": old_w, "المحجوبين (حالياً)": new_w,
                    "الحالة": status_text
                })
            else:
                counters["unchanged_fam"] += 1
                results.append({
                    "التسلسل": seq, "رقم البطاقة": card,
                    "الاسم الرباعي (سابقاً)": old_name, "الاسم الرباعي (حالياً)": new_name,
                    "الكلية (سابقاً)": old_t, "الكلية (حالياً)": new_t,
                    "المستحقة (سابقاً)": old_e, "المستحقة (حالياً)": new_e,
                    "المحجوبين (سابقاً)": old_w, "المحجوبين (حالياً)": new_w,
                    "الحالة": "✅ متطابق (بدون تغيير)" + audit_note
                })
                
        # الحالة الثانية: عائلة تم حذفها أو نقلها بالكامل من كشف الوكيل
        elif card in old_data:
            old_v = old_data[card]
            old_t, old_e, old_w = safe_ints(old_v)
            old_name = str(old_v.get("name", "")).strip()
            seq = old_v.get("seq", "-")
            
            counters["deleted_fam"] += 1
            counters["net_total"] -= old_t
            counters["net_eligible"] -= old_e
            counters["net_withheld"] -= old_w
            
            results.append({
                "التسلسل": seq, "رقم البطاقة": card,
                "الاسم الرباعي (سابقاً)": old_name, "الاسم الرباعي (حالياً)": "❌ (محذوف / منقول)",
                "الكلية (سابقاً)": old_t, "الكلية (حالياً)": 0,
                "المستحقة (سابقاً)": old_e, "المستحقة (حالياً)": 0,
                "المحجوبين (سابقاً)": old_w, "المحجوبين (حالياً)": 0,
                "الحالة": "🔴 محذوف من الوجبة"
            })
            
        # الحالة الثالثة: عائلة جديدة تم إضافتها بالكامل لكشف الوكيل الحالي
        elif card in new_data:
            new_v = new_data[card]
            new_t, new_e, new_w = safe_ints(new_v)
            new_name = str(new_v.get("name", "")).strip()
            seq = new_v.get("seq", "-")
            
            audit_note = " [⚠️ خلل حسابي بالملف الحالي]" if new_t != (new_e + new_w) else ""
            if audit_note: counters["data_errors"] += 1
            
            counters["added_fam"] += 1
            counters["net_total"] += new_t
            counters["net_eligible"] += new_e
            counters["net_withheld"] += new_w
            
            results.append({
                "التسلسل": seq, "رقم البطاقة": card,
                "الاسم الرباعي (سابقاً)": "✨ (مضاف حديثاً)", "الاسم الرباعي (حالياً)": new_name,
                "الكلية (سابقاً)": 0, "الكلية (حالياً)": new_t,
                "المستحقة (سابقاً)": 0, "المستحقة (حالياً)": new_e,
                "المحجوبين (سابقاً)": 0, "المحجوبين (حالياً)": new_w,
                "الحالة": "🟢 قيد مضاف جديد" + audit_note
            })
            
    return results, counters

# -----------------------------------------------------------------------------
# 5. محرك تقرير الـ Word للمتغير الحالي فقط مع التسلسل
# -----------------------------------------------------------------------------
def create_current_only_word_report(df, file_title):
    doc = Document()
    heading = doc.add_heading(f"كشف البيانات الحالية للوكيل: {file_title}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    report_rows = []
    for _, row in df.iterrows():
        quad_name = row.get("الاسم الرباعي (حالياً)", "-")
        if "❌" in str(quad_name):
            quad_name = row.get("الاسم الرباعي (سابقاً)", "-")
            
        report_rows.append({
            "التسلسل": row.get("التسلسل", "-"),
            "رقم البطاقة": row.get("رقم البطاقة", "-"),
            "الاسم الرباعي الحالي": quad_name,
            "الكلية الحالية": row.get("الكلية (حالياً)", 0),
            "المستحقة الحالية": row.get("المستحقة (حالياً)", 0),
            "المحجوبين حالياً": row.get("المحجوبين (حالياً)", 0),
            "حالة القيد": row.get("الحالة", "🔹 قيد تمويني")
        })
        
    report_df = pd.DataFrame(report_rows)
    cols = list(report_df.columns)[::-1]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = str(col)
        
    for _, r_row in report_df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = str(r_row[col])
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# 6. واجهة التطبيق التفاعلية والتحكم
# -----------------------------------------------------------------------------
tab1, tab2 = st.tabs(["🔎 تدقيق بيانات الوكيل الكاملة", "📜 الأرشيف التاريخي"])

with tab1:
    st.markdown("<h3 style='text-align: right;'>المنظومة الرقمية لعرض وجرد بيانات الوكيل بالكامل</h3>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1: new_file = st.file_uploader("الملف الجديد (الشهر الحالي)", type=['docx'], key="n_f")
    with col2: old_file = st.file_uploader("الملف القديم (الشهر السابق)", type=['docx'], key="o_f")

    col_btn1, col_btn2 = st.columns([3, 1])
    with col_btn1:
        run_match = st.button("🚀 جلب وعرض كافّة بيانات الوكيل")
    with col_btn2:
        clear_mem = st.button("🗑️ تفريغ الذاكرة لبدء وكيل جديد")
        if clear_mem:
            browser_memory["df"] = None
            browser_memory["counters"] = None
            browser_memory["filename"] = None
            for k in ['c_results', 'c_counters', 'c_filename']:
                if k in st.session_state: del st.session_state[k]
            st.rerun()

    if run_match:
        if old_file and new_file:
            with st.spinner('جاري قراءة الملفات وجرد الأسماء الرباعية بالتسلسلات...'):
                try:
                    old_data = extract_clean_records(old_file)
                    new_data = extract_clean_records(new_file)
                    results, counters = compare_records(old_data, new_data)
                    
                    if results:
                        st.session_state['c_results'] = pd.DataFrame(results).sort_values(by="التسلسل")
                        st.session_state['c_counters'] = counters
                        st.session_state['c_filename'] = new_file.name.rsplit('.', 1)[0]
                        
                        browser_memory["df"] = st.session_state['c_results']
                        browser_memory["counters"] = st.session_state['c_counters']
                        browser_memory["filename"] = st.session_state['c_filename']
                        
                        st.success("✅ تم جلب السجلات وتأمينها بالمتصفح ضد التحديث والمسح تلقائياً.")
                    else:
                        st.info("لم يتم العثور على سجلات صالحة.")
                except Exception as e:
                    st.error(f"حدث خطأ أثناء القراءة: {str(e)}")
        else:
            st.warning("يرجى تزويد النظام بملفات الـ Word لوكيلك أولاً.")

    # -------------------------------------------------------------------------
    # 7. لوحة العرض والفرز الإحصائي المتطورة
    # -------------------------------------------------------------------------
    if 'c_results' in st.session_state:
        df_res = st.session_state['c_results']
        cnt = st.session_state['c_counters']
        
        net_total = cnt.get('net_total', 0)
        added_fam = cnt.get('added_fam', 0)
        deleted_fam = cnt.get('deleted_fam', 0)
        unchanged_fam = cnt.get('unchanged_fam', 0)
        modified_fam = cnt.get('modified_fam', 0)
        data_errors = cnt.get('data_errors', 0)
        
        c1, c2, c3 = st.columns(3)
        with c1: st.markdown(f"<div class='report-box'>👥 إجمالي عوائل الوكيل: {len(df_res)}</div>", unsafe_allow_html=True)
        with c2: st.markdown(f"<div class='report-box'>📈 صافي تغيير الحصص المستحقة: {net_total:+d}</div>", unsafe_allow_html=True)
        with c3: 
            err_str = f" | ⚠️ خلل حسابي: {data_errors}" if data_errors > 0 else ""
            st.markdown(f"<div class='report-box'>✨ جديد: {added_fam} | 🟡 معدل: {modified_fam} | ❌ محذوف: {deleted_fam}{err_str}</div>", unsafe_allow_html=True)
        
        st.markdown("### 🔐 ترحيل سحابي دائم")
        if st.button("💾 حفظ بيانات هذا الوكيل بالكامل في Google Sheets"):
            if conn is not None:
                with st.spinner('جاري الإرسال للسحابة...'):
                    try:
                        try: existing_df = pd.DataFrame(conn.read(worksheet="Sheet1", ttl=0))
                        except Exception: existing_df = pd.DataFrame()
                        
                        df_to_save = df_res.copy()
                        df_to_save["وقت الحفظ"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        df_to_save["اسم الوكيل/الملف"] = st.session_state['c_filename']
                        
                        final_archive = pd.concat([existing_df, df_to_save], ignore_index=True)
                        conn.update(worksheet="Sheet1", data=final_archive)
                        st.success("🌟 تم حفظ الكشف كاملاً داخل ملف قوقل شيت بنجاح.")
                    except Exception as ex:
                        st.error(f"فشل الاتصال: {str(ex)}")
        
        st.markdown("---")
        st.markdown("<h4 style='text-align: right;'>🎯 لوحة البحث والفرز الذكي بأسماء الوكيل الحالية</h4>", unsafe_allow_html=True)
        
        filter_options = [
            f"📋 عرض كافّة سجلات الوكيل بالكامل ({len(df_res)})",
            f"✅ العوائل المستقرة والمتطابقة فقط ({unchanged_fam})",
            f"🟡 العوائل التي تغيرت حصصها التموينية أو أسماؤها ({modified_fam})",
            f"✨ العوائل المضافة حديثاً للوكيل ({added_fam})",
            f"❌ العوائل المحذوفة أو المنقولة من الوكيل ({deleted_fam})"
        ]
        selected_filter = st.selectbox("🎯 اختر الفئة المراد تدقيقها للوكيل:", filter_options)
        
        search_q = st.text_input("👤 اكتب الاسم الرباعي للمواطن أو رقم البطاقة للبحث الفوري:", "")
        
        has_status_col = "الحالة" in df_res.columns
        
        if has_status_col and "✅" in selected_filter:
            filtered_df = df_res[df_res["الحالة"].str.contains("متطابق", na=False)]
        elif has_status_col and "🟡" in selected_filter:
            filtered_df = df_res[df_res["الحالة"].str.contains("معدل", na=False)]
        elif has_status_col and "✨" in selected_filter:
            filtered_df = df_res[df_res["الحالة"].str.contains("مضاف", na=False)]
        elif has_status_col and "❌" in selected_filter:
            filtered_df = df_res[df_res["الحالة"].str.contains("محذوف", na=False)]
        else:
            filtered_df = df_res

        if search_q:
            filtered_df = filtered_df[
                filtered_df["الاسم الرباعي (حالياً)"].astype(str).str.contains(search_q, na=False) | 
                filtered_df["الاسم الرباعي (سابقاً)"].astype(str).str.contains(search_q, na=False) |
                filtered_df["رقم البطاقة"].astype(str).str.contains(search_q, na=False)
            ]
            
        st.markdown(f"<p style='text-align: right; color: gray;'>عدد النتائج الحالية: {len(filtered_df)} مواطن</p>", unsafe_allow_html=True)
        
        # 👤 العرض الديناميكي المرتّب للبطاقات
        for idx, row in filtered_df.iterrows():
            name_now = str(row.get("الاسم الرباعي (حالياً)", ""))
            name_old = str(row.get("الاسم الرباعي (سابقاً)", ""))
            card_num = str(row.get("رقم البطاقة", ""))
            seq_num = str(row.get("التسلسل", "-"))
            status_badge = row.get("الحالة", "🔹 قيد تمويني")
            
            display_name = name_old if "❌" in name_now else name_now
            
            card_class = "family-card-wrapper"
            if "متطابق" in status_badge: card_class += " card-matched"
            elif "معدل" in status_badge: card_class += " card-modified"
            elif "مضاف" in status_badge: card_class += " card-added"
            elif "محذوف" in status_badge: card_class += " card-deleted"
            
            box_title = f"🏷️ ت: {seq_num} | {display_name} | 📄 بطاقة: {card_num}"
            
            st.markdown(f"<div class='{card_class}'>", unsafe_allow_html=True)
            with st.expander(box_title):
                st.markdown(f"<p style='text-align:right;'><b>تقرير التدقيق البرمجي:</b> {status_badge}</p>", unsafe_allow_html=True)
                
                col_old_side, col_new_side = st.columns(2)
                
                with col_old_side:
                    st.markdown("<p style='text-align: center; font-weight: bold; color: #757575; border-bottom: 2px solid #e0e0e0; padding-bottom: 5px;'>📅 الحصص التموينية السابقة</p>", unsafe_allow_html=True)
                    m1, m2, m3 = st.columns(3)
                    with m1:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>👥 الكلية</span><span class='metric-value color-total'>{row.get('الكلية (سابقاً)', 0)}</span></div>", unsafe_allow_html=True)
                    with m2:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>✅ المستحقة</span><span class='metric-value color-eligible'>{row.get('المستحقة (سابقاً)', 0)}</span></div>", unsafe_allow_html=True)
                    with m3:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>🚫 المحجوبين</span><span class='metric-value color-withheld'>{row.get('المحجوبين (سابقاً)', 0)}</span></div>", unsafe_allow_html=True)
                        
                with col_new_side:
                    st.markdown("<p style='text-align: center; font-weight: bold; color: #1B5E20; border-bottom: 2px solid #a5d6a7; padding-bottom: 5px;'>🌟 الحصص التموينية الحالية (النهائية)</p>", unsafe_allow_html=True)
                    n1, n2, n3 = st.columns(3)
                    with n1:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>👥 الكلية</span><span class='metric-value color-total'>{row.get('الكلية (حالياً)', 0)}</span></div>", unsafe_allow_html=True)
                    with n2:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>✅ المستحقة</span><span class='metric-value color-eligible'>{row.get('المستحقة (حالياً)', 0)}</span></div>", unsafe_allow_html=True)
                    with n3:
                        st.markdown(f"<div class='metric-grid-box'><span class='metric-title'>🚫 المحجوبين</span><span class='metric-value color-withheld'>{row.get('المحجوبين (حالياً)', 0)}</span></div>", unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
        
        st.markdown("---")
        word_data = create_current_only_word_report(df_res, st.session_state['c_filename'])
        st.download_button(
            label="📥 تحميل كشف الحصص الحالية للوكيل (Word)", 
            data=word_data, 
            file_name=f"كشف_بيانات_الوكيل_الحالية_{st.session_state['c_filename']}.docx"
        )

with tab2:
    if st.button("🔄 جلب وتحديث الأرشيف المركزي السحابي"):
        try:
            if conn is not None:
                st.dataframe(conn.read(worksheet="Sheet1", ttl=0), use_container_width=True, hide_index=True)
            else:
                st.error("الاتصال بقاعدة البيانات السحابية مقطوع.")
        except Exception as ex: 
            st.error(ex)
